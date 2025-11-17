import io

from docx import Document
from reportlab.pdfgen import canvas


class ExportService:
    def generate_markdown(self, po, ba, technical_ba):
        markdown_lines = ["# Product Owner"]
        markdown_lines.extend([f"- {item}" for item in po])

        markdown_lines.append("\n# Business Analyst")
        markdown_lines.extend([f"- {item}" for item in ba])

        markdown_lines.append("\n# Technical BA")
        markdown_lines.extend([f"- {item}" for item in technical_ba])

        return "\n".join(markdown_lines)

    def generate_docx(self, po, ba, technical_ba):
        document = Document()

        document.add_heading("Product Owner", level=1)
        for item in po:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Business Analyst", level=1)
        for item in ba:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Technical BA", level=1)
        for item in technical_ba:
            document.add_paragraph(item, style="List Bullet")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_pdf(self, po, ba, technical_ba):
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer)

        y_position = 800

        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(50, y_position, "Product Owner")
        y_position -= 20

        pdf.setFont("Helvetica", 12)
        for item in po:
            pdf.drawString(70, y_position, f"- {item}")
            y_position -= 15

        y_position -= 10
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(50, y_position, "Business Analyst")
        y_position -= 20

        pdf.setFont("Helvetica", 12)
        for item in ba:
            pdf.drawString(70, y_position, f"- {item}")
            y_position -= 15

        y_position -= 10
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(50, y_position, "Technical BA")
        y_position -= 20

        pdf.setFont("Helvetica", 12)
        for item in technical_ba:
            pdf.drawString(70, y_position, f"- {item}")
            y_position -= 15

        pdf.save()
        buffer.seek(0)
        return buffer
